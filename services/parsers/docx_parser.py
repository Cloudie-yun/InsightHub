from pathlib import Path


def _dependency_error(package_name):
    return {
        "code": "missing_dependency",
        "message": f"Required parser dependency '{package_name}' is not installed.",
        "details": {"package": package_name},
    }


def _segment_sort_key(segment):
    return (
        segment.get("source_index", 0),
        segment.get("block_index", 0),
        segment.get("paragraph_index", 0),
        segment.get("segment_id", ""),
    )


def _iter_block_items(document):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document), "paragraph"
        elif isinstance(child, CT_Tbl):
            yield Table(child, document), "table"


def parse_docx(file_path):
    try:
        from docx import Document
    except Exception:
        return {
            "segments": [],
            "metadata": {},
            "errors": [_dependency_error("python-docx")],
        }

    file_path = Path(file_path)
    segments = []
    metadata = {
        "source_path": str(file_path),
        "paragraph_count": 0,
        "table_count": 0,
        "block_count": 0,
    }

    try:
        document = Document(file_path)
        block_index = 0
        paragraph_index = 0
        table_index = 0

        for block, block_type in _iter_block_items(document):
            block_index += 1

            if block_type == "paragraph":
                paragraph_index += 1
                text = (block.text or "").strip()
                if not text:
                    continue
                segments.append(
                    {
                        "segment_id": f"docx-block-{block_index}",
                        "source_type": "page",
                        "source_index": 1,
                        "block_index": block_index,
                        "paragraph_index": paragraph_index,
                        "text": text,
                        "heading": block.style.name if block.style else None,
                        "metadata": {
                            "block_type": "paragraph",
                            "char_count": len(text),
                            "parser_adapter": "docx",
                        },
                    }
                )
                continue

            table_index += 1
            for row_index, row in enumerate(block.rows, start=1):
                row_text_parts = []
                for cell in row.cells:
                    cell_text = (cell.text or "").strip()
                    if cell_text:
                        row_text_parts.append(cell_text)
                row_text = " | ".join(row_text_parts).strip()
                if not row_text:
                    continue
                segments.append(
                    {
                        "segment_id": f"docx-block-{block_index}-row-{row_index}",
                        "source_type": "page",
                        "source_index": 1,
                        "block_index": block_index,
                        "paragraph_index": row_index,
                        "text": row_text,
                        "metadata": {
                            "block_type": "table_row",
                            "table_index": table_index,
                            "row_index": row_index,
                            "char_count": len(row_text),
                            "parser_adapter": "docx",
                        },
                    }
                )

        metadata["paragraph_count"] = len(document.paragraphs)
        metadata["table_count"] = len(document.tables)
        metadata["block_count"] = block_index
        segments.sort(key=_segment_sort_key)
        return {"segments": segments, "metadata": metadata, "errors": []}
    except Exception as exc:
        return {
            "segments": [],
            "metadata": metadata,
            "errors": [
                {
                    "code": "docx_parse_error",
                    "message": "Unable to parse DOCX file.",
                    "details": {"exception": str(exc), "source_path": str(file_path)},
                }
            ],
        }
